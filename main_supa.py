import os
import tempfile
from typing import List, Dict
import re
import numpy as np
import fitz  # PyMuPDF
import markdown
import requests
from dotenv import load_dotenv
from fastapi import FastAPI, File, UploadFile, Request, Form, Query
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from openai import OpenAI
from prompts import prompts
from docx import Document

load_dotenv()
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")
TABLE_NAME = "financials"

headers = {
    "apikey": SUPABASE_SERVICE_KEY,
    "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
    "Content-Type": "application/json",
}

openai_key = os.getenv("OPENAI_KEY")
app = FastAPI()
os.makedirs("static", exist_ok=True)
os.makedirs("templates", exist_ok=True)
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")
client = OpenAI(api_key=openai_key)

COMPLETION_MODEL = "gpt-4o-mini"
TOP_K_CHUNKS = 10  # Number of top chunks to use
uploaded_pdf_paths: List[str] = []

def markdown_to_html(md_text: str) -> str:
    return markdown.markdown(md_text, extensions=["tables"])

def extract_pdf_text(pdf_path: str) -> List[Dict]:
    try:
        doc = fitz.open(pdf_path)
        pages_text = []
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            text = page.get_text()
            pages_text.append({"page_number": page_num + 1, "text": text.strip()})
        doc.close()
        return pages_text
    except Exception as e:
        print(f"Error extracting PDF text from {pdf_path}: {e}")
        return []

def clean_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s\.,;:!?()-]', ' ', text)
    return text.strip()

def chunk_text(text: str, max_chunk_size: int = 1000, overlap: int = 100) -> List[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + max_chunk_size
        chunk = text[start:end]
        chunks.append(chunk.strip())
        start += max_chunk_size - overlap
    return chunks

def get_embedding(text: str) -> np.ndarray:
    try:
        response = client.embeddings.create(input=[text], model="text-embedding-3-small")
        return np.array(response.data[0].embedding, dtype=np.float32)
    except Exception as e:
        print(f"Embedding error: {e}")
        return np.zeros(1536)

def cosine_similarity(vec1: np.ndarray, vec2: np.ndarray) -> float:
    dot_product = np.dot(vec1, vec2)
    norm1 = np.linalg.norm(vec1)
    norm2 = np.linalg.norm(vec2)
    if norm1 == 0 or norm2 == 0:
        return 0.0
    return dot_product / (norm1 * norm2)

def clean_latex(text: str) -> str:
    return text.replace("\\[", "").replace("\\]", "").replace("\\(", "").replace("\\)", "").replace("$$", "").replace("\\text{", "").replace("}", "").strip()

def analyze_documents_enhanced(pdf_paths: List[str], question: str, file_names: List[str] = None):
    steps = ["📄 Extracting and chunking text from all PDFs..."]
    all_chunks = []
    
    for i, path in enumerate(pdf_paths):
        pages = extract_pdf_text(path)
        full_text = " ".join([page["text"] for page in pages])
        cleaned_text = clean_text(full_text)
        text_chunks = chunk_text(cleaned_text)

        for idx, chunk in enumerate(text_chunks):
            embedding = get_embedding(chunk)
            all_chunks.append({
                "text": chunk,
                "chunk_number": idx + 1,
                "source": f"Document_{i+1}",
                "embedding": embedding
            })

    steps.append("🧠 Calculating similarity scores...")
    try:
        query_embedding = get_embedding(question)
        for chunk in all_chunks:
            chunk["similarity_score"] = float(cosine_similarity(query_embedding, chunk["embedding"]))
        relevant_chunks = sorted(all_chunks, key=lambda x: x["similarity_score"], reverse=True)[:TOP_K_CHUNKS]
    except Exception as e:
        print(f"Similarity error: {e}")
        relevant_chunks = all_chunks[:TOP_K_CHUNKS]

    merged_text = "\n\n".join([chunk["text"] for chunk in relevant_chunks])

    prompt = f"""You are an expert assistant helping answer questions from financial documents. Use only the information provided below to answer the question.

Context:
{merged_text}

Question:
{question}

Answer concisely, citing key facts, figures, and chunk numbers if possible."""

    response = client.chat.completions.create(
        model=COMPLETION_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=300,
    )

    final_answer = clean_latex(response.choices[0].message.content.strip())

    if "|" in final_answer and "-" in final_answer:
        final_answer = markdown_to_html(final_answer)

    return {
        "steps": steps,
        "answer": final_answer,
        "chunks_used": len(relevant_chunks),
        "total_chunks": len(all_chunks)
    }

@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    return templates.TemplateResponse("index.html", {"request": request, "prompts": prompts})

@app.post("/upload")
async def upload_files(files: List[UploadFile] = File(...)):
    for old_path in uploaded_pdf_paths:
        try:
            os.unlink(old_path)
        except Exception:
            pass
    uploaded_pdf_paths.clear()
    file_names = []
    for file in files:
        if not file.filename.lower().endswith(".pdf"):
            return JSONResponse(status_code=400, content={"error": f"Invalid file type: {file.filename}"})
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        content = await file.read()
        with open(temp_file.name, "wb") as f:
            f.write(content)
        uploaded_pdf_paths.append(temp_file.name)
        file_names.append(file.filename)
    return {"filenames": file_names, "status": "Files uploaded successfully", "file_count": len(files)}

@app.post("/analyze")
async def analyze(prompt_key: str = Form(...), custom_query: str = Form(None)):
    if not uploaded_pdf_paths:
        return JSONResponse(status_code=400, content={"error": "No documents uploaded yet"})
    question = custom_query.strip() if custom_query else prompts.get(prompt_key)
    if not question:
        return JSONResponse(status_code=400, content={"error": "Invalid or missing query"})
    file_names = [f"Document_{i+1}" for i in range(len(uploaded_pdf_paths))]
    result = analyze_documents_enhanced(uploaded_pdf_paths, question, file_names)
    return {
        "answer": result["answer"],
        "steps": result["steps"],
        "chunks_used": result["chunks_used"],
        "total_chunks": result["total_chunks"]
    }

@app.post("/analyze_custom")
async def analyze_custom(custom_query: str = Form(...)):
    if not uploaded_pdf_paths:
        return JSONResponse(status_code=400, content={"error": "No documents uploaded yet"})
    if not custom_query.strip():
        return JSONResponse(status_code=400, content={"error": "Custom query cannot be empty"})
    file_names = [f"Document_{i+1}" for i in range(len(uploaded_pdf_paths))]
    result = analyze_documents_enhanced(uploaded_pdf_paths, custom_query.strip(), file_names)
    return {
        "answer": result["answer"],
        "steps": result["steps"],
        "chunks_used": result["chunks_used"],
        "total_chunks": result["total_chunks"]
    }

@app.get("/scrape_nse")
async def scrape_nse(tickers: str = Query(..., description="Comma separated tickers")):
    ticker_list = [t.strip().upper() for t in tickers.split(",") if t.strip()]
    if not ticker_list:
        return JSONResponse(status_code=400, content={"error": "No valid tickers provided"})
    in_clause = ",".join(ticker_list)
    supabase_query_url = f"{SUPABASE_URL}/rest/v1/{TABLE_NAME}?ticker=in.({in_clause})"
    response = requests.get(supabase_query_url, headers=headers)
    if response.status_code == 200:
        return response.json()
    else:
        return JSONResponse(status_code=response.status_code, content={"error": "Failed to fetch from Supabase", "details": response.text})



@app.post("/generate_report")
async def generate_report():
    if not uploaded_pdf_paths:
        return JSONResponse(status_code=400, content={"error": "No documents uploaded yet"})
    
    # Get the first two prompts from the prompts dict (ordered by insertion)
    keys = list(prompts.keys())[:3]  # e.g. ["Business", "Competition"]
    file_names = [f"Document_{i+1}" for i in range(len(uploaded_pdf_paths))]
    
    doc = Document()
    doc.add_heading("Full Equity Research Report", level=1)
    
    for key in keys:
        prompt_text = prompts.get(key)
        if not prompt_text:
            continue  # skip if prompt not found
        
        # Call your analysis function for each prompt
        result = analyze_documents_enhanced(uploaded_pdf_paths, prompt_text, file_names)
        
        # Add a heading for the prompt section
        doc.add_heading(key, level=2)
        # Add the text answer for that section
        doc.add_paragraph(result["answer"])
    
    temp_doc_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(temp_doc_path)
    
    return FileResponse(
        path=temp_doc_path,
        filename="full_equity_report.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )